from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import pandas as pd
from docx import Document
import io
import os
import logging
from collections import defaultdict

# --- Setup Logging ---
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("NEXAI-API")

app = FastAPI(title="NEXAI Timetable Generator")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Advanced Error Handling ---
@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    logger.error(f"Unhandled error: {str(exc)}")
    return JSONResponse(
        status_code=500,
        content={"detail": "Internal Server Error", "message": str(exc)}
    )

class TimetableGenerator:
    def __init__(self):
        self.days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
        self.room_bookings = defaultdict(lambda: defaultdict(set))
        self.teacher_bookings = defaultdict(lambda: defaultdict(set))
        self.section_bookings = defaultdict(lambda: defaultdict(set))

    def parse_excel(self, file_content: bytes):
        try:
            excel_file = pd.ExcelFile(io.BytesIO(file_content), engine='openpyxl')
            sheet_names = excel_file.sheet_names
            
            # Helper to find sheets regardless of case
            find_sheet = lambda name: next((s for s in sheet_names if s.lower() == name.lower()), None)
            
            s_teacher = find_sheet('Teacher')
            s_sections = find_sheet('Sections')
            s_rooms = find_sheet('rooms')

            if not all([s_teacher, s_sections, s_rooms]):
                missing = [n for n, s in zip(['Teacher', 'Sections', 'rooms'], [s_teacher, s_sections, s_rooms]) if not s]
                raise ValueError(f"Missing sheets: {', '.join(missing)}")

            # Load and clean headers
            df_teacher = pd.read_excel(excel_file, sheet_name=s_teacher)
            df_sections = pd.read_excel(excel_file, sheet_name=s_sections)
            df_rooms = pd.read_excel(excel_file, sheet_name=s_rooms)

            df_teacher.columns = [str(c).strip() for c in df_teacher.columns]
            df_sections.columns = [str(c).strip() for c in df_sections.columns]
            df_rooms.columns = [str(c).strip() for c in df_rooms.columns]

            return {'teacher': df_teacher, 'sections': df_sections, 'rooms': df_rooms}
        except Exception as e:
            raise ValueError(f"Excel Structure Error: {str(e)}")

    def generate_timetables(self, data):
        timetables = defaultdict(list)
        all_rooms = data['rooms'].to_dict('records')
        
        teacher_map = {}
        for idx, row in data['teacher'].iterrows():
            # Robust column name checking for 'Name' or 'Nmae'
            name = str(row.get('Name', row.get('Nmae', ''))).strip()
            courses = str(row.get('courses', '')).split(',')
            try:
                ch = int(row.get('credit hours', 1))
            except:
                ch = 1
                
            if not name or name == 'nan':
                continue # Skip empty rows
                
            for c in courses:
                teacher_map[c.strip()] = {"name": name, "credit_hours": ch}

        for _, row in data['sections'].iterrows():
            section = str(row.get('Section', '')).strip()
            subjects = [s.strip() for s in str(row.get('Subject', '')).split(',') if s.strip()]
            
            if not section or not subjects:
                continue

            for sub in subjects:
                t_data = teacher_map.get(sub)
                if not t_data:
                    raise ValueError(f"Teacher not found for subject: '{sub}' in Section: '{section}'")

                teacher = t_data["name"]
                is_lab = sub.lower().endswith('lab')
                duration = 3 if is_lab else t_data["credit_hours"]
                display_sub_name = f"{sub} (lab)" if is_lab and "lab" not in sub.lower() else sub

                scheduled = False
                for day in self.days:
                    if scheduled: break
                    
                    start_search = 9 if teacher.lower().endswith('main') else 8
                    for start_h in range(start_search, 16 - duration + 1):
                        if any(12 <= h < (14 if day == 'Friday' else 13) for h in range(start_h, start_h + duration)):
                            continue

                        slots = [f"{h}:00" for h in range(start_h, start_h + duration)]
                        found_room = None
                        for r in all_rooms:
                            r_id = str(r.get('room id', ''))
                            r_type = str(r.get('type', '')).lower()
                            
                            if is_lab != ('lab' in r_type): continue

                            if all(r_id not in self.room_bookings[day][s] and 
                                   teacher not in self.teacher_bookings[day][s] and
                                   section not in self.section_bookings[day][s] for s in slots):
                                found_room = r_id
                                break
                        
                        if found_room:
                            for s in slots:
                                self.room_bookings[day][s].add(found_room)
                                self.teacher_bookings[day][s].add(teacher)
                                self.section_bookings[day][s].add(section)
                            
                            timetables[section].append({
                                'day': day, 'time': f"{start_h}:00", 
                                'end_time': f"{start_h + duration}:00",
                                'subject': display_sub_name, 'teacher': teacher, 
                                'room': f"[{found_room}]"
                            })
                            scheduled = True
                            break
                
                if not scheduled:
                    raise ValueError(f"Could not find a slot for {sub} (Section {section}) due to conflicts.")

        return timetables

    def generate_word_doc(self, timetables):
        doc = Document()
        for section, entries in sorted(timetables.items()):
            doc.add_heading(f'Section: {section}', level=1)
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Time'
            for i, day in enumerate(self.days): hdr_cells[i+1].text = day

            lookup = defaultdict(str)
            for e in entries:
                start_h = int(e['time'].split(':')[0])
                end_h = int(e['end_time'].split(':')[0])
                for h in range(start_h, end_h):
                    lookup[(e['day'], f"{h}:00")] = f"{e['subject']}\n({e['teacher']})\n{e['room']}"

            for h in range(8, 16):
                row_cells = table.add_row().cells
                row_cells[0].text = f"{h}:00 - {h+1}:00"
                for i, day in enumerate(self.days):
                    if h == 12: row_cells[i+1].text = "BREAK"
                    elif day == 'Friday' and h == 13: row_cells[i+1].text = "JUMMAH"
                    else: row_cells[i+1].text = lookup[(day, f"{h}:00")]
            doc.add_page_break()
        
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        return out

@app.post("/download-timetable")
async def handle_download(file: UploadFile = File(...)):
    try:
        content = await file.read()
        gen = TimetableGenerator()
        data = gen.parse_excel(content)
        timetables = gen.generate_timetables(data)
        doc_io = gen.generate_word_doc(timetables)
        
        return StreamingResponse(
            doc_io, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=timetable.docx"}
        )
    except ValueError as ve:
        logger.warning(f"Validation Error: {str(ve)}")
        raise HTTPException(status_code=400, detail=str(ve))
    except Exception as e:
        logger.error(f"Unexpected processing error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")

@app.get("/")
async def root():
    return {"status": "active", "service": "NEXAI Timetable Generator"}

# No uvicorn.run block needed for Vercel